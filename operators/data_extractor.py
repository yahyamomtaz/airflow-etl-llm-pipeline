import logging
import os
from docx import Document
from docx.oxml.shared import qn
import re
from typing import Any
from utils.storage import get_storage, LocalStorage

log = logging.getLogger(__name__)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_DIR = os.path.join(BASE_DIR, "..", "data", "descriptions")
DOCX_DIR = os.path.abspath(DOCX_DIR)

FIELD_MAP = {
    "Autore": "author",
    "Autore secondario": "second_author", 
    "Titolo": "title",
    "Pubblicazione": "publication",
    "Dimensioni": "dimensions",
    "Peso": "weight",
    "Spessore dei fogli": "thickness",
    "Collocazione": "location",
    "Segnatura": "signature",
    "Impronta": "imprint",
    "Disposizione del testo": "text_layout",
    "Righe": "lines",
    "Richiami": "requests",
    "Legatura": "binding",
    "Lingua": "language_info",
    "Nomi significativi": "significant_names",
    "Stato di conservazione": "condition_info",
    "Decorazione": "decoration",
    "Descrizione fisica": "physical_description",
    "": None  # fallback
}

# Mapping from numeric filename prefix (e.g. "01") → (collection_id, collection_name)
# Add entries here to map specific prefixes to collections in your institution.
# Any prefix not listed here falls back to PREFIX_COLLECTION_DEFAULT.
PREFIX_COLLECTION_MAP = {
}

# Default collection for simple-format documents — configure for your institution.
PREFIX_COLLECTION_DEFAULT = (1, "Default Collection")


def normalize_simple_book_title(title):
    """Return the canonical short title for simple-format books."""
    if not title:
        return ""

    normalized = re.sub(r'\s+', ' ', title).strip()
    normalized = re.sub(
        r'^scheda\s+descrittiva\s+',
        '',
        normalized,
        flags=re.IGNORECASE,
    ).strip()
    return re.split(r'\s*[:–-]\s*', normalized, maxsplit=1)[0].strip()

def strip_html_tags(text):
    """Remove HTML tags from text while preserving the content"""
    if not text:
        return text
    clean_text = re.sub(r'<[^>]+>', '', text)
    return clean_text.strip().rstrip(".")

def is_simple_format_filename(filename):
    """Return True if the filename follows the new simple format: NN-<title>.docx"""
    return bool(re.match(r'^\s*\d{2}\s*[^0-9A-Za-z]+\s*', filename, re.IGNORECASE))


def determine_collection_from_filename(filename):
    """Determine which collection this file belongs to based on the number"""
    # Extract number from filename
    # Updated regex to handle book numbers with parentheses like 5D43(1)
    match = re.search(r'Scheda descrittiva_([A-Za-z0-9()]+)_VERIFICATA', filename, re.IGNORECASE)
    if not match:
        log.warning("Could not extract book number from filename: %s", filename)
        return None  # Return None instead of tuple
    
    book_number = match.group(1).upper()
    
    # Collection A: book numbers starting with 5 — configure name/id for your institution
    if book_number.startswith('5'):
        collection_id = 4
        collection_name = "Collection A"
        log.info("Found collection: %s for file: %s", collection_name, book_number)
        return collection_id, collection_name, book_number

    # Collection B: book numbers starting with 4 — configure name/id for your institution
    elif book_number.startswith('4'):
        collection_id = 3
        collection_name = "Collection B"
        log.info("Found collection: %s for file: %s", collection_name, book_number)
        return collection_id, collection_name, book_number
    
    # Unknown collection - still return the book number but with unknown collection
    else:
        log.warning("Unknown collection for book: %s", book_number)
        return None, "Unknown", book_number


def extract_simple_format_docx(docx_path):
    """Extract data from a simple-format docx file (NN-<title>.docx).
    
    These files contain only two fields:
      - Titolo (title)
      - Descrizione (description / physical_description)

    The collection and book identifier are derived from the filename prefix.
    Because there is no author, a title-based slug is used instead.
    """
    filename = os.path.basename(docx_path)

    # Derive the numeric prefix (e.g. "01" from "01-scheda descrittiva Compactio I.docx")
    prefix_match = re.match(r'^\s*(\d{2})\s*[^0-9A-Za-z]+\s*', filename)
    if not prefix_match:
        log.warning("Could not extract numeric prefix from filename: %s", filename)
        return {}
    prefix = prefix_match.group(1)

    # Look up collection from PREFIX_COLLECTION_MAP, falling back to default
    collection_id, collection_name = PREFIX_COLLECTION_MAP.get(prefix, PREFIX_COLLECTION_DEFAULT)

    doc = Document(docx_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    title = ""
    description = ""

    for i, para in enumerate(paragraphs):
        # Title line: starts with "Titolo" (case-insensitive)
        title_match = re.match(r'^Titolo[:\s]+(.+)$', para, re.IGNORECASE)
        if title_match:
            title = title_match.group(1).strip()
            continue

        # Description line: starts with "Descrizione" or "Descriz"
        desc_match = re.match(r'^Descrizi?o?n?e?[:\s]+(.*)$', para, re.IGNORECASE)
        if desc_match:
            desc_text = desc_match.group(1).strip()
            # Description may continue on subsequent lines
            remaining = paragraphs[i + 1:]
            parts = [desc_text] if desc_text else []
            for next_para in remaining:
                # Stop if we hit another field label
                if re.match(r'^[A-ZÀ-ÿ][a-zÀ-ÿ ]+:', next_para):
                    break
                parts.append(next_para)
            description = ' '.join(parts).strip()
            continue

    # If title not found via label, fall back to the first non-empty paragraph
    if not title and paragraphs:
        title = paragraphs[0]

    # book_number uses only the short title (before any " – " subtitle).
    # title (the full Titolo value) is stored as-is in book_descriptions.title.
    # book_slug is the URL-safe version used in paths/manifest URLs (e.g. "compactio-v").
    short_title = normalize_simple_book_title(title)
    book_number = short_title  # e.g. "Compactio V"
    book_slug = re.sub(r'[^a-z0-9]+', '-', short_title.lower()).strip('-')  # e.g. "compactio-v"

    # Simple-format books have no author: keep author_slug empty.
    # The database_operator, manifest generator, and viewers will fall back
    # to using just book_slug as the slug.
    author_slug = ""

    # Image folder name matches the books/ folder on disk (e.g. "01-Compactio I")
    image_folder_name = f"{prefix}-{short_title}"

    data = {
        "title": title,
        "description": description,
        "book_number": book_number,
        "book_slug": book_slug,
        "collection_id": collection_id,
        "collection_name": collection_name,
        "author_slug": author_slug,
        "author": "",
        # Explicit image folder name so the manifest generator can find the images
        "image_folder_name": image_folder_name,
    }

    log.info(
        "Simple format: title=%r, collection=%s (id=%s), book_number=%r, slug=%r",
        title, collection_name, collection_id, book_number, book_slug,
    )
    log.info("Image folder: %r", image_folder_name)
    log.debug("Extracted data: %s", data)
    return data

def extract_data_from_docx(docx_path, storage=None):
    """Dispatch to the correct extractor based on filename format.

    Downloads from S3 to a temp file automatically when the storage
    backend is S3; transparent to callers.
    """
    if storage is None:
        storage = get_storage()

    filename = os.path.basename(docx_path)
    tmp_path = None
    local_path = docx_path

    # If not on local storage, download to a temporary file first
    if not isinstance(storage, LocalStorage):
        tmp_path = storage.local_path(docx_path)
        local_path = tmp_path

    try:
        if is_simple_format_filename(filename):
            log.info("Detected simple format (NN-prefix): %s", filename)
            return extract_simple_format_docx(local_path)
        return _extract_standard_docx(local_path)
    finally:
        if tmp_path:
            try:
                os.unlink(tmp_path)
            except OSError:
                pass


def _extract_standard_docx(docx_path):
    doc = Document(docx_path)
    
    # Extract text with hyperlinks preserved
    full_text_with_links = ""
    paragraph_texts = []
    
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            # Try to extract with hyperlinks
            text_with_links = extract_hyperlinks_from_paragraph(paragraph)
            paragraph_texts.append(text_with_links)
            full_text_with_links += text_with_links + "\n"
    
    # Also get plain text for regex matching (fallback)
    plain_text = "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])
    
    data = {}
    for key, db_col in FIELD_MAP.items():
        if not db_col:
            continue
        
        # First try to find the field in text with links
        match = re.search(rf"{re.escape(key)}:\s*(.+?)(?=\n[A-ZÀ-ÿ][a-zÀ-ÿ ]+?:|\Z)", full_text_with_links, re.DOTALL | re.IGNORECASE)
        
        # If not found, try plain text
        if not match:
            match = re.search(rf"{re.escape(key)}:\s*(.+?)(?=\n[A-ZÀ-ÿ][a-zÀ-ÿ ]+?:|\Z)", plain_text, re.DOTALL | re.IGNORECASE)
        
        if match:
            value = match.group(1).strip()
            # Clean up whitespace but preserve HTML tags
            value = re.sub(r'\s+', ' ', value)
            # Clean up any malformed HTML
            value = re.sub(r'<a\s+href="([^"]*)"[^>]*>\s*</a>', '', value)  # Remove empty links
            data[db_col] = value

    # Handle author slug
    author = data.get("author", "")
    if author:
        author_clean = strip_html_tags(author)
        author_slug = re.sub(r'[^a-z0-9]+', '-', author_clean.lower()).strip('-')
    else:
        author_slug = "unknown-author"
        log.warning("No author found, using default slug")
    
    # Determine collection metadata
    collection_info = determine_collection_from_filename(os.path.basename(docx_path))
    if collection_info is None:
        stem = os.path.splitext(os.path.basename(docx_path))[0]
        data["book_number"] = stem
        data["collection_id"] = None
        data["collection_name"] = "Unknown"
        data["author_slug"] = author_slug
        log.warning(
            "Arbitrary filename — using stem as book_number=%r, collection=Unknown", stem
        )
    elif isinstance(collection_info, tuple) and len(collection_info) == 3:
        collection_id, collection_name, book_number = collection_info
        data["book_number"] = book_number
        data["collection_id"] = collection_id
        data["collection_name"] = collection_name
        data["author_slug"] = author_slug
    else:
        log.warning("Unexpected collection_info format: %s", collection_info)

    log.debug("Extracted data: %s", data)
            
    return data

def extract_hyperlinks_from_paragraph(paragraph):
    """Extract paragraph text while preserving hyperlinks as HTML.
    
    This handles multiple Word hyperlink structures:
    1. Hyperlinks as run parents (standard)
    2. Hyperlinks in paragraph XML (some Word versions)
    """
    parts = []
    
    # Method 1: Check each run's parent for hyperlink
    for run in paragraph.runs:
        run_parent = run._element.getparent()
        if run_parent is not None and run_parent.tag is not None and run_parent.tag.endswith('hyperlink'):
            rel_id = run_parent.get(qn('r:id'))
            hyperlink_url = None
            if rel_id and rel_id in paragraph.part.rels:
                hyperlink_url = paragraph.part.rels[rel_id].target_ref

            run_text = run.text or ''
            if hyperlink_url and run_text.strip():
                parts.append(f'<a href="{hyperlink_url}" target="_blank">{run_text}</a>')
            else:
                parts.append(run_text)
        else:
            parts.append(run.text or '')

    joined = ''.join(parts)

    # Method 2: If no hyperlinks found in runs, check paragraph XML directly
    if '<a href=' not in joined:
        try:
            # Get paragraph XML
            para_xml = paragraph._element.xml
            if isinstance(para_xml, bytes):
                para_xml = para_xml.decode('utf-8')
            else:
                para_xml = str(para_xml)
            
            # Check if there are hyperlinks in the XML
            if '<w:hyperlink' in para_xml or 'hyperlink' in para_xml.lower():
                # Extract hyperlink elements with their rIds
                import re
                
                # Find all w:hyperlink elements with their r:id
                hyperlink_pattern = r'<w:hyperlink[^>]*r:id="([^"]+)"[^>]*>(.*?)</w:hyperlink>'
                hyperlink_matches = re.findall(hyperlink_pattern, para_xml, re.DOTALL)
                
                if hyperlink_matches:
                    # We found hyperlinks in the XML, need to reconstruct text with links
                    result_parts = []
                    last_pos = 0
                    
                    # Get plain text to know the structure
                    plain_text = paragraph.text
                    
                    for rel_id, hyperlink_content in hyperlink_matches:
                        # Extract the text from the hyperlink content
                        text_matches = re.findall(r'<w:t[^>]*>([^<]+)</w:t>', hyperlink_content)
                        link_text = ''.join(text_matches)
                        
                        # Try to get the URL from relationships
                        try:
                            if rel_id in paragraph.part.rels:
                                url = paragraph.part.rels[rel_id].target_ref
                                
                                # Find where this text appears in the plain text
                                if link_text in plain_text[last_pos:]:
                                    pos = plain_text.find(link_text, last_pos)
                                    # Add text before the link
                                    if pos > last_pos:
                                        result_parts.append(plain_text[last_pos:pos])
                                    # Add the link
                                    result_parts.append(f'<a href="{url}" target="_blank">{link_text}</a>')
                                    last_pos = pos + len(link_text)
                        except:
                            # If we can't resolve the URL, just add the text
                            result_parts.append(link_text)
                    
                    # Add any remaining text
                    if last_pos < len(plain_text):
                        result_parts.append(plain_text[last_pos:])
                    
                    joined = ''.join(result_parts)
        
        except Exception as e:
            # If XML parsing fails, fall back to plain text
            pass

    # If there were no hyperlinks, fall back to the simple paragraph text
    if '<a href=' not in joined:
        return (paragraph.text or '').strip()

    return joined

from airflow.models import BaseOperator


class DataExtractorOperator(BaseOperator):
    """
    Airflow operator that extracts data from Word files.
    
    Pulls new_books from sensor XCom, extracts data, and pushes books_data to XCom.
    """
    
    template_fields = ('descriptions_dir',)
    
    def __init__(
        self,
        descriptions_dir: str = '/opt/airflow/data/descriptions',
        sensor_task_id: str = 'wait_for_new_books',
        **kwargs
    ):
        super().__init__(**kwargs)
        self.descriptions_dir = descriptions_dir
        self.sensor_task_id = sensor_task_id
    
    def execute(self, context):
        from operators.processed_books import update_book_status

        self.log.info("Extracting Data from Word Files")

        # Get new books from sensor
        new_books = context['ti'].xcom_pull(key='new_books', task_ids=self.sensor_task_id)

        extracted_data = []
        failed_extractions = []

        if not new_books:
            self.log.warning("No new books detected by sensor — scanning all .docx files")
            docx_files = [f for f in os.listdir(self.descriptions_dir)
                          if f.lower().endswith('.docx') and not f.startswith('~')]
            if not docx_files:
                raise FileNotFoundError(f"No .docx files found in {self.descriptions_dir}")

            self.log.info("Found %d Word document(s) (scanning all)", len(docx_files))
            for filename in docx_files:
                file_path = os.path.join(self.descriptions_dir, filename)
                self.log.info("Processing: %s", filename)
                try:
                    data = extract_data_from_docx(file_path)
                    if data:
                        missing = self._validate_required_fields(data)
                        if missing:
                            self.log.warning("Missing required fields %s in %s", missing, filename)
                            failed_extractions.append({'file': filename, 'reason': f'Missing fields: {missing}', 'data': data})
                        else:
                            extracted_data.append(data)
                            self.log.info("Extracted: %s", data.get('book_number'))
                    else:
                        self.log.error("No data returned from extraction for %s", filename)
                        failed_extractions.append({'file': filename, 'reason': 'No data returned from extraction'})
                except Exception as e:
                    self.log.error("Error extracting %s: %s", filename, e)
                    failed_extractions.append({'file': filename, 'reason': str(e)})
        else:
            self.log.info("Processing %d new book(s) from sensor", len(new_books))
            for book_info in new_books:
                file_path = os.path.join(self.descriptions_dir, book_info['docx_file'])
                self.log.info("Processing: %s (Book: %s)", book_info['docx_file'], book_info['book_number'])
                try:
                    data = extract_data_from_docx(file_path)
                    if data:
                        missing = self._validate_required_fields(data)
                        if missing:
                            self.log.warning("Missing required fields %s", missing)
                            failed_extractions.append({
                                'file': book_info['docx_file'],
                                'book_number': book_info['book_number'],
                                'reason': f'Missing fields: {missing}',
                                'data': data
                            })
                            update_book_status(book_info['book_number'], 'extracted', False, f'Missing fields: {missing}')
                        else:
                            extracted_data.append(data)
                            self.log.info("Extracted: %s", data.get('book_number'))
                            update_book_status(data.get('book_number'), 'extracted', True)
                    else:
                        self.log.error("No data returned from extraction for %s", book_info['docx_file'])
                        failed_extractions.append({
                            'file': book_info['docx_file'],
                            'book_number': book_info['book_number'],
                            'reason': 'No data returned from extraction'
                        })
                        update_book_status(book_info['book_number'], 'extracted', False, 'No data returned from extraction')
                except Exception as e:
                    self.log.error("Error extracting %s: %s", book_info['docx_file'], e)
                    failed_extractions.append({
                        'file': book_info['docx_file'],
                        'book_number': book_info['book_number'],
                        'reason': str(e)
                    })
                    update_book_status(book_info['book_number'], 'extracted', False, str(e))

        if failed_extractions:
            self.log.warning("%d book(s) failed extraction", len(failed_extractions))
            for fail in failed_extractions:
                self.log.warning("  - %s: %s", fail.get('book_number', fail.get('file')), fail['reason'])
            context['task_instance'].xcom_push(key='failed_extractions', value=failed_extractions)

        context['task_instance'].xcom_push(key='books_data', value=extracted_data)
        self.log.info("Extracted %d books successfully, %d failed", len(extracted_data), len(failed_extractions))
        return len(extracted_data)
    
    def _validate_required_fields(self, data):
        """Validate that required fields are present"""
        missing = []
        if not data.get('book_number'): missing.append('book_number')
        if not data.get('collection_name'): missing.append('collection_name')
        # Simple-format books (identified by 'image_folder_name') have no author by design.
        is_simple_format = bool(data.get('image_folder_name'))
        if not is_simple_format and not data.get('author_slug'):
            missing.append('author_slug')
        return missing


if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)
    for filename in os.listdir(DOCX_DIR):
        if filename.lower().endswith(".docx"):
            file_path = os.path.join(DOCX_DIR, filename)
            log.info("--- Extracting from %s ---", filename)
            extract_data_from_docx(file_path)
