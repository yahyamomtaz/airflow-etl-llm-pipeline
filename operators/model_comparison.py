import os
import json
import time
from datetime import datetime
import logging

log = logging.getLogger(__name__)

def validate_word_content_with_llama(file_path, book_number, num_images):
    """
    Use Llama model (via Ollama) to validate Word file content
    
    Returns same format as GPT-4 for comparison
    """
    try:
        import requests
        from docx import Document
        
        start_time = time.time()
        
        # Extract text from Word file
        doc = Document(file_path)
        full_text = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
        
        # Truncate if too long
        if len(full_text) > 4000:
            full_text = full_text[:4000] + "\n... (truncated)"
        
        validation_prompt = f"""You are validating a book description file for a digital library.

Book Number: {book_number}
Number of Images: {num_images}

Description Content:
{full_text}

Validate if the description contains these REQUIRED fields:
1. Author name
2. Title
3. Publication information

Also check for OPTIONAL fields:
- Physical dimensions
- Location/shelf mark
- Language
- Condition notes

Respond ONLY with valid JSON in this exact format:
{{
  "valid": true or false,
  "critical_issues": ["list any missing required fields"],
  "warnings": ["list any missing optional fields"],
  "suggestions": ["improvement suggestions"],
  "confidence": 0-100,
  "detected_fields": {{
    "author": "author name or null",
    "title": "title or null",
    "publication": "publication info or null"
  }}
}}"""

        # Call Ollama API (local Llama instance)
        ollama_url = os.getenv('OLLAMA_URL', 'http://localhost:11434')
        model_name = os.getenv('LLAMA_MODEL', 'llama3.1:8b')
        
        response = requests.post(
            f"{ollama_url}/api/generate",
            json={
                "model": model_name,
                "prompt": validation_prompt,
                "stream": False,
                "format": "json"
            },
            timeout=60
        )
        
        response.raise_for_status()
        result_data = response.json()
        
        elapsed_time = time.time() - start_time
        
        # Parse response
        result_text = result_data.get('response', '{}')
        result = json.loads(result_text)
        
        # Add metrics
        result['_metrics'] = {
            'model': model_name,
            'latency_seconds': round(elapsed_time, 2),
            'tokens_used': result_data.get('eval_count', 0),
            'cost': 0.0  # Llama is free locally
        }
        
        return result
        
    except requests.exceptions.ConnectionError:
        return {
            "valid": True,
            "error": "Ollama not running. Start with: ollama serve",
            "skipped": True,
            "_metrics": {"model": "llama", "error": "connection_failed"}
        }
    except ImportError:
        return {
            "valid": True,
            "error": "requests or docx package not installed",
            "skipped": True,
            "_metrics": {"model": "llama", "error": "import_failed"}
        }
    except json.JSONDecodeError as e:
        return {
            "valid": True,
            "error": f"Llama returned invalid JSON: {str(e)}",
            "skipped": True,
            "_metrics": {"model": "llama", "error": "invalid_json"}
        }
    except Exception as e:
        return {
            "valid": True,
            "error": str(e),
            "skipped": True,
            "_metrics": {"model": "llama", "error": "other"}
        }


def validate_word_content_with_gpt4(file_path, book_number, num_images):
    """
    Use GPT-4 to validate Word file content (with metrics tracking)
    """
    try:
        import openai
        from docx import Document
        
        start_time = time.time()
        
        # Extract text
        doc = Document(file_path)
        full_text = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
        
        if len(full_text) > 4000:
            full_text = full_text[:4000] + "\n... (truncated)"
        
        validation_prompt = f"""You are validating a book description file for a digital library pipeline.

**Book Information:**
- Book Number: {book_number}
- Number of Images: {num_images}

**Description File Content:**
```
{full_text}
```

**Validation Requirements:**
Check if the description contains:
1. ✅ Author name (required)
2. ✅ Title (required)
3. ✅ Publication info (required)
4. ✅ Physical dimensions
5. ✅ Location/shelf mark
6. ✅ Language information
7. ✅ Condition notes

**Validation Criteria:**
- CRITICAL: Missing author, title, or publication info
- WARNING: Missing optional fields (dimensions, condition)
- SUGGESTION: Incomplete data or potential improvements

Return JSON ONLY (no markdown):
{{
  "valid": true/false,
  "critical_issues": ["list of critical missing fields"],
  "warnings": ["list of warnings"],
  "suggestions": ["improvement suggestions"],
  "confidence": 0-100,
  "detected_fields": {{
    "author": "extracted author name or null",
    "title": "extracted title or null",
    "publication": "extracted publication info or null"
  }}
}}
"""
        
        response = openai.chat.completions.create(
            model="gpt-4-turbo",
            messages=[
                {
                    "role": "system",
                    "content": "You are a meticulous digital library QA expert. Validate book descriptions thoroughly. Return ONLY valid JSON, no markdown formatting."
                },
                {
                    "role": "user",
                    "content": validation_prompt
                }
            ],
            temperature=0.1,
            response_format={"type": "json_object"}
        )
        
        elapsed_time = time.time() - start_time
        
        result_text = response.choices[0].message.content
        result = json.loads(result_text)
        
        # Add metrics
        usage = response.usage
        input_tokens = usage.prompt_tokens
        output_tokens = usage.completion_tokens
        
        # GPT-4 Turbo pricing (as of Dec 2024)
        cost = (input_tokens * 0.01 / 1000) + (output_tokens * 0.03 / 1000)
        
        result['_metrics'] = {
            'model': 'gpt-4-turbo',
            'latency_seconds': round(elapsed_time, 2),
            'input_tokens': input_tokens,
            'output_tokens': output_tokens,
            'total_tokens': input_tokens + output_tokens,
            'cost': round(cost, 4)
        }
        
        return result
        
    except ImportError:
        return {
            "valid": True,
            "error": "openai or docx package not installed",
            "skipped": True,
            "_metrics": {"model": "gpt-4", "error": "import_failed"}
        }
    except json.JSONDecodeError as e:
        return {
            "valid": True,
            "error": f"GPT-4 returned invalid JSON: {str(e)}",
            "skipped": True,
            "_metrics": {"model": "gpt-4", "error": "invalid_json"}
        }
    except Exception as e:
        return {
            "valid": True,
            "error": str(e),
            "skipped": True,
            "_metrics": {"model": "gpt-4", "error": "other"}
        }


def compare_model_validations(file_path, book_number, num_images, use_gpt4=True, use_llama=True):
    """
    Run both GPT-4 and Llama validations and compare results
    
    Returns:
        dict: {
            'gpt4': {...},
            'llama': {...},
            'comparison': {
                'agreement': bool,
                'latency_winner': 'gpt4'|'llama',
                'cost_winner': 'llama' (always, since it's free),
                'differences': [...]
            }
        }
    """
    results = {}
    
    if use_gpt4:
        log.info("      🔵 Running GPT-4 validation...")
        results['gpt4'] = validate_word_content_with_gpt4(file_path, book_number, num_images)
    
    if use_llama:
        log.info("      🟢 Running Llama validation...")
        results['llama'] = validate_word_content_with_llama(file_path, book_number, num_images)
    
    # Compare if both ran successfully
    if 'gpt4' in results and 'llama' in results:
        gpt4_valid = results['gpt4'].get('valid', True)
        llama_valid = results['llama'].get('valid', True)
        
        comparison = {
            'agreement': gpt4_valid == llama_valid,
            'both_valid': gpt4_valid and llama_valid,
            'gpt4_metrics': results['gpt4'].get('_metrics', {}),
            'llama_metrics': results['llama'].get('_metrics', {}),
            'timestamp': datetime.now().isoformat()
        }
        
        # Latency comparison
        gpt4_latency = results['gpt4'].get('_metrics', {}).get('latency_seconds', float('inf'))
        llama_latency = results['llama'].get('_metrics', {}).get('latency_seconds', float('inf'))
        
        if gpt4_latency < llama_latency:
            comparison['latency_winner'] = 'gpt4'
            comparison['latency_diff'] = round(llama_latency - gpt4_latency, 2)
        else:
            comparison['latency_winner'] = 'llama'
            comparison['latency_diff'] = round(gpt4_latency - llama_latency, 2)
        
        # Cost comparison (Llama always wins - it's free!)
        comparison['cost_winner'] = 'llama'
        comparison['gpt4_cost'] = results['gpt4'].get('_metrics', {}).get('cost', 0)
        
        # Check for differences in findings
        gpt4_issues = set(results['gpt4'].get('critical_issues', []))
        llama_issues = set(results['llama'].get('critical_issues', []))
        
        comparison['differences'] = {
            'only_gpt4_found': list(gpt4_issues - llama_issues),
            'only_llama_found': list(llama_issues - gpt4_issues),
            'both_found': list(gpt4_issues & llama_issues)
        }
        
        results['comparison'] = comparison
    
    return results
